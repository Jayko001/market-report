from docx import Document
from data_processing import get_dataframe
from main import get_company_info
from sqlalchemy import create_engine
import json

df = get_dataframe()


def generate_document(company_name, competitors_info, output_file):
    doc = Document()
    doc.add_heading(f'Company Competitor Analysis for {company_name}', 0)
    
    for competitor, info in competitors_info.items():
        doc.add_heading(competitor, level=1)
        doc.add_heading('About:', level=2)
        doc.add_paragraph(info['about'])
        doc.add_heading('Customers:', level=2)
        doc.add_paragraph(info['customers'])
        doc.add_heading('Pricing:', level=2)
        doc.add_paragraph(info['pricing'])

    doc.save(output_file)
    print(f"Document saved as {output_file}")

def extract_company_names(df):
    # Get a list of all the company names
    company_names = df['companies'].unique().tolist()
    return company_names

def main():
    try:
        company_name = "Perceive Now"
        competitor_names = extract_company_names(df)
        print(competitor_names)
        # competitor_names = ['IPwe','Cardinal Intellectual Property', 'Amplified', 'YouScan']
        competitors_info = get_company_info(company_name, competitor_names)
        output_file = f'Competitor_Analysis_for_{company_name}.docx'
        # print(competitors_info)
        generate_document(company_name, competitors_info, output_file)

    except Exception as e:
        print(f"An error occurred: {e}")

if __name__ == "__main__":
    main()

